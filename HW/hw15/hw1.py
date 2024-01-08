from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()
doc.add_paragraph("Hello Python")
doc.save("hello_python.docx")

read_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in read_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
print("Жирный текст:", bold_text)

new_doc = Document()
new_paragraph = new_doc.add_paragraph("Новый абзац с измененным шрифтом и размером.")
new_run = new_paragraph.runs[0]
new_run.bold = True
new_run.font.size = Pt(14)
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
new_doc.save("new_file.docx")